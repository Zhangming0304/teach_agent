"""
文件解析模块 - 支持多格式文件解析
支持格式：图片(jpg/png/webp/bmp)、PDF、Word(.docx)、纯文本(.txt)
"""
import os
import uuid
from dataclasses import dataclass, field

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 支持的文件扩展名
IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.webp', '.bmp'}
PDF_EXTENSIONS = {'.pdf'}
DOCX_EXTENSIONS = {'.docx'}
TEXT_EXTENSIONS = {'.txt'}

# 扫描件判定阈值：页面文字少于此字数视为扫描件
SCAN_TEXT_THRESHOLD = 20


@dataclass
class ParseResult:
    """文件解析结果"""
    mode: str = "vision"          # "vision" 或 "text"
    images: list = field(default_factory=list)   # mode=vision 时的图片路径列表
    text: str = ""                # mode=text 时的文字内容
    file_type: str = "image"      # "image" / "pdf_scan" / "pdf_text" / "docx" / "text"
    page_count: int = 0           # 页数


def _get_extension(filepath: str) -> str:
    """获取文件扩展名（小写，含点号）"""
    _, ext = os.path.splitext(filepath)
    return ext.lower()


def _parse_image(filepath: str) -> ParseResult:
    """
    解析图片文件 — 直接返回图片路径，走多模态视觉识别

    Args:
        filepath: 图片文件路径

    Returns:
        ParseResult(mode="vision", images=[filepath], file_type="image", page_count=1)
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"图片文件不存在: {filepath}")

    return ParseResult(
        mode="vision",
        images=[filepath],
        text="",
        file_type="image",
        page_count=1
    )


def _parse_pdf(filepath: str) -> ParseResult:
    """
    解析 PDF 文件 — 区分扫描件和文字件

    - 扫描件（页面文字少于阈值字数）→ 导出为图片，走视觉识别
    - 文字 PDF → 提取文字内容，走纯文本批改

    Args:
        filepath: PDF 文件路径

    Returns:
        ParseResult
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"PDF 文件不存在: {filepath}")

    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("需要安装 PyMuPDF 库: pip install PyMuPDF")

    doc = fitz.open(filepath)
    page_count = len(doc)

    if page_count == 0:
        doc.close()
        return ParseResult(
            mode="text",
            images=[],
            text="",
            file_type="pdf_text",
            page_count=0
        )

    # 检测是否为扫描件：逐页检查文字量
    is_scan = True
    all_text_parts = []

    for page in doc:
        page_text = page.get_text("text").strip()
        all_text_parts.append(page_text)
        if len(page_text) >= SCAN_TEXT_THRESHOLD:
            is_scan = False

    if is_scan:
        # 扫描件 → 每页导出为图片
        images = []
        for page_idx in range(page_count):
            page = doc[page_idx]
            # 高分辨率渲染 (2x)
            mat = fitz.Matrix(2, 2)
            pix = page.get_pixmap(matrix=mat)

            img_filename = f"pdf_page_{uuid.uuid4().hex}_{page_idx + 1}.png"
            img_path = os.path.join(UPLOAD_DIR, img_filename)
            pix.save(img_path)
            images.append(img_path)

        doc.close()
        return ParseResult(
            mode="vision",
            images=images,
            text="",
            file_type="pdf_scan",
            page_count=page_count
        )
    else:
        # 文字 PDF → 提取全部文字
        full_text = "\n\n".join(
            f"--- 第 {i + 1} 页 ---\n{text}"
            for i, text in enumerate(all_text_parts)
            if text
        )
        doc.close()
        return ParseResult(
            mode="text",
            images=[],
            text=full_text,
            file_type="pdf_text",
            page_count=page_count
        )


def _parse_docx(filepath: str) -> ParseResult:
    """
    解析 Word 文件 — 提取文字内容，走纯文本批改

    Args:
        filepath: Word 文件路径

    Returns:
        ParseResult(mode="text", text=..., file_type="docx")
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"Word 文件不存在: {filepath}")

    try:
        from docx import Document
    except ImportError:
        raise ImportError("需要安装 python-docx 库: pip install python-docx")

    doc = Document(filepath)

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 也尝试提取表格中的文字
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    full_text = "\n".join(paragraphs)

    return ParseResult(
        mode="text",
        images=[],
        text=full_text,
        file_type="docx",
        page_count=1
    )


def _parse_text(filepath: str) -> ParseResult:
    """
    解析纯文本文件 — 直接读取

    Args:
        filepath: 文本文件路径

    Returns:
        ParseResult(mode="text", text=..., file_type="text")
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"文本文件不存在: {filepath}")

    # 尝试多种编码读取
    content = None
    for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                content = f.read()
            break
        except (UnicodeDecodeError, LookupError):
            continue

    if content is None:
        raise ValueError(f"无法识别文件编码: {filepath}")

    return ParseResult(
        mode="text",
        images=[],
        text=content.strip(),
        file_type="text",
        page_count=1
    )


def parse_file(filepath: str) -> ParseResult:
    """
    解析单个文件，根据扩展名自动选择解析方式

    Args:
        filepath: 文件路径

    Returns:
        ParseResult 解析结果

    Raises:
        ValueError: 不支持的文件格式
        FileNotFoundError: 文件不存在
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"文件不存在: {filepath}")

    ext = _get_extension(filepath)

    if ext in IMAGE_EXTENSIONS:
        return _parse_image(filepath)
    elif ext in PDF_EXTENSIONS:
        return _parse_pdf(filepath)
    elif ext in DOCX_EXTENSIONS:
        return _parse_docx(filepath)
    elif ext in TEXT_EXTENSIONS:
        return _parse_text(filepath)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，支持的格式: 图片(jpg/png/webp/bmp)、PDF、Word(.docx)、纯文本(.txt)")


def parse_files(filepaths: list[str]) -> ParseResult:
    """
    解析多个文件并合并结果

    合并策略：
    - 如果所有文件都是视觉模式 → 合并图片列表
    - 如果所有文件都是文本模式 → 合并文本内容
    - 如果混合模式 → 优先使用视觉模式（图片文件走视觉，文本文件的内容也会附带）

    Args:
        filepaths: 文件路径列表

    Returns:
        合并后的 ParseResult
    """
    if not filepaths:
        raise ValueError("文件列表为空")

    results = []
    for fp in filepaths:
        try:
            result = parse_file(fp)
            results.append(result)
        except Exception as e:
            raise ValueError(f"解析文件 {os.path.basename(fp)} 失败: {str(e)}")

    if len(results) == 1:
        return results[0]

    # 检查模式：是否有 vision 和 text 混合
    has_vision = any(r.mode == "vision" for r in results)
    has_text = any(r.mode == "text" for r in results)

    all_images = []
    all_texts = []
    total_pages = 0
    file_types = set()

    for r in results:
        all_images.extend(r.images)
        if r.text:
            all_texts.append(r.text)
        total_pages += r.page_count
        file_types.add(r.file_type)

    # 确定合并后的 file_type
    if len(file_types) == 1:
        merged_file_type = file_types.pop()
    else:
        merged_file_type = "mixed"

    if has_vision and not has_text:
        # 纯视觉模式
        return ParseResult(
            mode="vision",
            images=all_images,
            text="",
            file_type=merged_file_type,
            page_count=total_pages
        )
    elif has_text and not has_vision:
        # 纯文本模式
        return ParseResult(
            mode="text",
            images=[],
            text="\n\n".join(all_texts),
            file_type=merged_file_type,
            page_count=total_pages
        )
    else:
        # 混合模式 → 优先视觉模式，文本内容附在 text 字段中
        return ParseResult(
            mode="vision",
            images=all_images,
            text="\n\n".join(all_texts),
            file_type=merged_file_type,
            page_count=total_pages
        )
